"""DOCX text extraction using python-docx."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from app.models.domain import ExtractedDocument
from app.processing.extractors.base import BaseExtractor


class DocxExtractor(BaseExtractor):
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        document = Document(BytesIO(content))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs).strip()
        pages = [(1, text)] if text else []
        return ExtractedDocument(filename=filename, text=text, pages=pages)
